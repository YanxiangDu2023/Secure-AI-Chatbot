import os
import logging
import pymupdf
from langchain.schema import Document
from docx import Document as DocxDocument
from fastapi import APIRouter, HTTPException, File, UploadFile
from app.utils import create_vector_db, get_llm
from langchain_text_splitters import RecursiveCharacterTextSplitter
from app.models import vector_db_state
from app.globals import chat_file_mapping
import aiofiles
from concurrent.futures import ThreadPoolExecutor
import asyncio
import pandas as pd
import magic

router = APIRouter()
process_executor = ThreadPoolExecutor(max_workers=4)
SUPPORTED_TYPES = {
    'application/pdf': 'pdf',
    'application/vnd.openxmlformats-officedocument.wordprocessingml.document': 'docx',
    'application/vnd.ms-excel': 'xls',
    'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet': 'xlsx',
    'text/csv': 'csv',
    'application/msword': 'doc',
    'application/vnd.ms-excel': 'xls',
    'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet.main+xml': 'xlsx',
}

llm = get_llm()
UPLOAD_FOLDER = "./uploads"
VECTOR_STORE_DIR = "./vector_store"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
shared_vector_db = None
os.makedirs(VECTOR_STORE_DIR, exist_ok=True)
EMBEDDING_MODEL = "nomic-embed-text"

def extract_text_pymupdf(pdf_path):
    """Extract text from a PDF using PyMuPDF (fitz)."""
    doc = pymupdf.open(pdf_path)
    text = "\n".join([page.get_text("text") for page in doc])
    if not text.strip():
        raise ValueError("No text found in the PDF.")
    return text

async def detect_file_type(file_path: str):
    """Detect file type using python-magic"""
    try:
        mime = magic.Magic(mime=True)
        detected_type = mime.from_file(file_path)
        file_type = SUPPORTED_TYPES.get(detected_type, 'unknown')
        
        # Fallback to file extension if MIME type fails
        if file_type == 'unknown':
            ext = os.path.splitext(file_path)[1].lower()
            extension_map = {
                '.docx': 'docx',
                '.xlsx': 'xlsx',
                '.xls': 'xls',
                '.csv': 'csv',
                '.pdf': 'pdf'
            }
            file_type = extension_map.get(ext, 'unknown')
            
        print(f"Detected type: {detected_type} -> {file_type}")  # Debug log
        return file_type
        
    except Exception as e:
        logging.error(f"Type detection error: {str(e)}")
        return 'unknown'

async def extract_text(file_path: str, file_type: str):
    """Unified text extraction for different file types"""
    if file_type == 'pdf':
        return await asyncio.get_event_loop().run_in_executor(
            process_executor,
            lambda: extract_text_pymupdf(file_path)
        )
    elif file_type == 'docx':
        return await asyncio.get_event_loop().run_in_executor(
            process_executor,
            lambda: '\n'.join([p.text for p in DocxDocument(file_path).paragraphs])
        )
    elif file_type in ('xlsx', 'xls'):
        return await process_excel(file_path, file_type)
    else:
        raise ValueError(f"Unsupported file type: {file_type}")

async def process_excel(file_path: str, file_type: str):
    """Dedicated Excel processor"""
    return await asyncio.get_event_loop().run_in_executor(
        process_executor,
        lambda: pd_read_excel(file_path, file_type)
    )

def pd_read_excel(file_path: str, file_type: str):
    """Excel-specific text conversion"""
    text = []
    try:
        xls = pd.ExcelFile(file_path)
        for sheet_name in xls.sheet_names:
            df = pd.read_excel(xls, sheet_name=sheet_name)
            text.append(f"\n\nSheet: {sheet_name}\n{df.to_string()}")
        return '\n'.join(text)
    except Exception as e:
        raise ValueError(f"Excel processing failed: {str(e)}")

async def process_file(chat_id: str, file_path: str, filename: str):
    """Process files for a chat session, aggregating documents"""
    try:
        existing_files = chat_file_mapping.get(chat_id, [])
        # Detect file type
        file_type = await detect_file_type(file_path)
        if file_type == 'unknown':
            print(f"Unsupported file: {filename}")
            raise ValueError("Unsupported file format")

        # Extract text
        text = await extract_text(file_path, file_type)
        print(f"File content extracted ({file_type.upper()})")

        # Create document and split
        document = await asyncio.get_event_loop().run_in_executor(
            process_executor,
            lambda: Document(page_content=text, metadata={"source": filename})
        )
        
        chunks = await asyncio.get_event_loop().run_in_executor(
            process_executor,
            lambda: RecursiveCharacterTextSplitter(
                chunk_size=1200, 
                chunk_overlap=200
            ).split_documents([document])
        )
        print(f"Split into {len(chunks)} chunks")

        # Create vector DB
        vector_db = await asyncio.get_event_loop().run_in_executor(
            process_executor,
            lambda: create_vector_db(chunks, chat_id=chat_id)
        )

        # Update state
        vector_db_state.set_vector_db(vector_db)
        chat_file_mapping.setdefault(chat_id, []).append(file_path)

        return True
    except Exception as e:
        logging.error(f"Processing error: {e}")
        raise

@router.post("/{chat_id}")
async def upload_file(chat_id: str, file: UploadFile = File(...)):
    """Upload endpoint with guaranteed completion before response"""
    try:
        # Create upload directory
        chat_upload_dir = os.path.join(UPLOAD_FOLDER, chat_id)
        os.makedirs(chat_upload_dir, exist_ok=True)
        file_path = os.path.join(chat_upload_dir, file.filename)

        # Async file save
        async with aiofiles.open(file_path, "wb") as f:
            await f.write(await file.read())

        # Process with async wrapper
        success = await process_file(chat_id, file_path, file.filename)
        
        if not success:
            raise HTTPException(500, "File processing failed")

        return {"message": f"File processed successfully."}

    except Exception as e:
        logging.error(f"Upload error: {e}")
        raise HTTPException(500, "File processing failed") from e